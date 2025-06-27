#!/usr/bin/env python3
"""
agent.py — Ikirōne Agent (Final Corrected Order)
"""
import os
import requests
import msal
import json
import urllib.parse
import shutil
import io
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime, timezone

# Document processing libraries
import docx
from pdfminer.high_level import extract_text
from PIL import Image
from PIL.ExifTags import TAGS

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Header
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

# LangChain Imports
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage
from langchain.agents import AgentExecutor, create_openai_tools_agent
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.tools import tool
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# --- CONFIG & CONSTANTS ---
BASE_DIR = Path(__file__).parent
STORE_DIR = BASE_DIR / "faiss_store"
CHAT_LOG_FILE = BASE_DIR / "chat_log.json"
load_dotenv(BASE_DIR / ".env")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")
IKIRONE_API_KEY = os.getenv("IKIRONE_API_KEY")
MS_CLIENT_ID = os.getenv("MS_CLIENT_ID")
MS_CLIENT_SECRET = os.getenv("MS_CLIENT_SECRET")
MS_TENANT_ID = os.getenv("MS_TENANT_ID")
GRAPH_USER_UPN = os.getenv("GRAPH_USER_UPN")

# --- FASTAPI APP & MODELS DEFINITION ---
# ** THE FIX IS HERE: Defining the app and models before they are used **
app = FastAPI(title="Ikirōne Agent")
class ChatRequest(BaseModel):
    message: str

# --- HELPER FUNCTIONS ---
graph_token_cache = {"token": None, "expires_at": 0}

def get_graph_token() -> str:
    now = datetime.now(timezone.utc).timestamp()
    if graph_token_cache["token"] and graph_token_cache["expires_at"] > now:
        return graph_token_cache["token"]
    auth_client = msal.ConfidentialClientApplication(
        MS_CLIENT_ID, authority=f"https://login.microsoftonline.com/{MS_TENANT_ID}",
        client_credential=MS_CLIENT_SECRET
    )
    result = auth_client.acquire_token_for_client(["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        graph_token_cache["token"] = result["access_token"]
        graph_token_cache["expires_at"] = now + result.get("expires_in", 3600) - 60
        return result["access_token"]
    raise RuntimeError(f"MSAL token error: {result.get('error')}")

def get_drive_items_recursively(token: str, item_id: str = "root", path: str = "") -> List[Dict[str, Any]]:
    all_items = []
    url = f"https://graph.microsoft.com/v1.0/users/{GRAPH_USER_UPN}/drive/items/{item_id}/children"
    while url:
        r = requests.get(url, headers={"Authorization": f"Bearer {token}"})
        r.raise_for_status()
        data = r.json()
        items = data.get("value", [])
        for item in items:
            current_path = os.path.join(path, item.get('name', ''))
            item['path'] = current_path
            if "folder" in item:
                print(f"  Entering folder: {current_path}")
                all_items.extend(get_drive_items_recursively(token, item['id'], path=current_path))
            elif "file" in item:
                all_items.append(item)
        url = data.get("@odata.nextLink")
    return all_items

def save_chat_history(history: List[BaseMessage]):
    with open(CHAT_LOG_FILE, 'w') as f:
        json.dump([message.dict() for message in history], f, indent=2)

def load_chat_history() -> List[BaseMessage]:
    if not CHAT_LOG_FILE.exists():
        return []
    with open(CHAT_LOG_FILE, 'r') as f:
        try:
            history_dicts = json.load(f)
            messages = []
            for msg_dict in history_dicts:
                if msg_dict.get('type') == 'human':
                    messages.append(HumanMessage(**msg_dict))
                elif msg_dict.get('type') == 'ai':
                    messages.append(AIMessage(**msg_dict))
            return messages
        except (json.JSONDecodeError, TypeError):
            return []

# --- AGENT TOOLS ---
@tool
def get_current_time() -> str:
    """Returns the current date and time."""
    return datetime.now().isoformat()

@tool
def search_recent_emails(top: int = 10) -> List[Dict[str, Any]]:
    """Fetches the absolute most recent emails (up to 10)."""
    token = get_graph_token()
    url = f"https://graph.microsoft.com/v1.0/users/{GRAPH_USER_UPN}/mailFolders/Inbox/messages"
    params = {"$select": "id,subject,bodyPreview,from,receivedDateTime,webLink", "$top": min(top, 10), "$orderby": "receivedDateTime desc"}
    r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, params=params)
    r.raise_for_status()
    return r.json().get("value", [])

class LocalIndexSearchInput(BaseModel):
    query: str = Field(..., description="The user's question for semantic search.")

@tool(args_schema=LocalIndexSearchInput)
def query_local_index(query: str) -> str:
    """Performs a deep, semantic search over the local index of emails and documents."""
    if not STORE_DIR.exists() or not any(STORE_DIR.iterdir()):
        return "The local index has not been built yet. Please ask the user for permission to build it."
    embedder = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    vector_store = FAISS.load_local(str(STORE_DIR), embedder, allow_dangerous_deserialization=True)
    retriever = vector_store.as_retriever(search_kwargs={"k": 20})
    docs = retriever.invoke(query)
    context = "\n\n---\n\n".join(
        f"SOURCE: {d.metadata.get('source', 'N/A')}\nNAME: {d.metadata.get('name', d.metadata.get('subject', 'N/A'))}\n"
        f"DATE: {d.metadata.get('date', 'N/A')}\nPREVIEW: {d.page_content[:300]}"
        for d in docs
    )
    return f"Found {len(docs)} relevant items. Summary:\n{context}"

@tool
def list_recent_onedrive_files(top: int = 10) -> List[Dict[str, Any]]:
    """Lists the most recent files from the user's OneDrive."""
    token = get_graph_token()
    url = f"https://graph.microsoft.com/v1.0/users/{GRAPH_USER_UPN}/drive/root/children"
    params = {"$top": min(top, 20), "$orderby": "lastModifiedDateTime desc"}
    r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, params=params)
    r.raise_for_status()
    return r.json().get("value", [])

@tool
def build_local_index() -> str:
    """Performs a deep ingestion of items to build the local index, including EXIF data from images."""
    if STORE_DIR.exists(): shutil.rmtree(STORE_DIR)
    STORE_DIR.mkdir(exist_ok=True)
    
    token = get_graph_token()
    all_docs = []
    
    print("Fetching emails...")
    email_url = f"https://graph.microsoft.com/v1.0/users/{GRAPH_USER_UPN}/mailFolders/Inbox/messages"
    email_params = {"$select": "subject,body,from,receivedDateTime", "$top": 200, "$orderby": "receivedDateTime desc"}
    emails = requests.get(email_url, headers={"Authorization": f"Bearer {token}"}, params=email_params).json().get("value", [])
    for email in emails:
        all_docs.append(Document(
            page_content=email.get('body', {}).get('content', ''),
            metadata={"source": "Email", "name": email.get('subject', 'N/A'), "sender": email.get('from', {}).get('emailAddress', {}).get('name', 'N/A'), "date": email.get('receivedDateTime', '')}
        ))

    target_folders = ["Documents", "Pictures"]
    all_files = []
    for folder_name in target_folders:
        print(f"Looking for '{folder_name}' folder...")
        try:
            folder_url = f"https://graph.microsoft.com/v1.0/users/{GRAPH_USER_UPN}/drive/root:/{folder_name}"
            folder_id = requests.get(folder_url, headers={"Authorization": f"Bearer {token}"}).json().get('id')
            if folder_id:
                print(f"Found '{folder_name}'. Traversing...")
                all_files.extend(get_drive_items_recursively(token, item_id=folder_id, path=folder_name))
        except Exception as e:
            print(f"  Could not find or process folder '{folder_name}'. Error: {e}")

    supported_text = ['.txt', '.md', '.docx', '.pdf']
    supported_images = ['.jpg', '.jpeg', '.png', '.heic', '.tiff']
    for item in all_files:
        file_name = item.get('name', 'Untitled')
        file_lower = file_name.lower()
        is_text = any(file_lower.endswith(ext) for ext in supported_text)
        is_image = any(file_lower.endswith(ext) for ext in supported_images)
        if not (is_text or is_image): continue

        print(f"  Processing file: {item.get('path', file_name)}")
        try:
            content_response = requests.get(f"https://graph.microsoft.com/v1.0/users/{GRAPH_USER_UPN}/drive/items/{item['id']}/content", headers={"Authorization": f"Bearer {token}"})
            content_response.raise_for_status()
            doc_content = ""
            if is_text:
                if file_name.endswith(('.txt', '.md')): doc_content = content_response.text
                elif file_name.endswith('.docx'): doc_content = "\n".join([p.text for p in docx.Document(io.BytesIO(content_response.content)).paragraphs])
                elif file_name.endswith('.pdf'): doc_content = extract_text(io.BytesIO(content_response.content))
            elif is_image:
                try:
                    img = Image.open(io.BytesIO(content_response.content))
                    exif_data = img._getexif()
                    if exif_data:
                        exif_text = f"Image metadata for {file_name}:\n"
                        for tag_id, value in exif_data.items():
                            tag_name = TAGS.get(tag_id, tag_id)
                            if tag_name in ['Make', 'Model', 'DateTimeOriginal', 'LensModel']:
                                exif_text += f" - {tag_name}: {value}\n"
                        doc_content = exif_text
                except Exception as exif_error: doc_content = f"Image file: {file_name}. No readable EXIF data. Error: {exif_error}"
            
            if doc_content:
                all_docs.append(Document(page_content=doc_content, metadata={"source": "OneDrive", "name": file_name, "path": item.get('path', ''), "date": item.get('lastModifiedDateTime', '')}))
        except Exception as e:
            print(f"    Could not process file {file_name}. Error: {e}")

    if not all_docs: return "Could not find any items to index."
    
    splits = RecursiveCharacterTextSplitter.from_tiktoken_encoder(chunk_size=1000, chunk_overlap=200).split_documents(all_docs)
    embedder = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    batch_size, vector_store = 200, None
    print(f"Embedding {len(splits)} document chunks in batches...")
    for i in range(0, len(splits), batch_size):
        batch = splits[i:i + batch_size]
        if vector_store is None: vector_store = FAISS.from_documents(batch, embedder)
        else: vector_store.add_documents(batch)
        print(f"  ...processed batch {i//batch_size + 1}")
    
    if vector_store: vector_store.save_local(str(STORE_DIR))
    return f"Successfully built the local index with {len(all_docs)} total items."

# --- AGENT SETUP ---
tools = [get_current_time, search_recent_emails, list_recent_onedrive_files, query_local_index, build_local_index]
prompt = ChatPromptTemplate.from_messages([
    ("system", """You are Ikirōne..."""),
    MessagesPlaceholder(variable_name="chat_history"),
    ("human", "{input}"),
    MessagesPlaceholder(variable_name="agent_scratchpad"),
])

llm = ChatOpenAI(model=OPENAI_MODEL, openai_api_key=OPENAI_API_KEY)
agent = create_openai_tools_agent(llm, tools, prompt)
agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)
chat_history: List[BaseMessage] = load_chat_history()
print(f"Loaded {len(chat_history)} messages from previous sessions.")

# --- API ENDPOINTS ---
@app.post("/chat")
async def chat(req: ChatRequest, x_api_key: str = Header(None)):
    global chat_history
    if x_api_key != IKIRONE_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API Key")
    response = await agent_executor.ainvoke({"input": req.message, "chat_history": chat_history})
    chat_history.append(HumanMessage(content=response["input"]))
    chat_history.append(AIMessage(content=response["output"]))
    save_chat_history(chat_history)
    return {"response": response["output"]}

@app.get("/")
async def read_index(): return FileResponse(BASE_DIR / 'static' / 'index.html')
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("agent:app", host="0.0.0.0", port=8000, reload=True)
